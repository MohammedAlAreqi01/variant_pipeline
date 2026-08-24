#!/usr/bin/env python3
"""
Combine classified variant tables + ancestry stats across all genes and
gnomAD versions into the formatted summary tables (ENG/ACVRL1/combined/
v4/PALB2 structure, sensitivity/specificity/PPV/NPV, etc.) for the thesis.

Called by the `build_tables` Snakemake rule. Uses python-docx - swap in
your existing table-formatting code from the thesis here.
"""
import argparse
import pandas as pd
from docx import Document


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--classified", nargs="+", required=True)
    p.add_argument("--stats", nargs="+", required=True)
    p.add_argument("--output", required=True)
    args = p.parse_args()

    classified_dfs = [pd.read_csv(f, sep="\t") for f in args.classified]
    stats_dfs = [pd.read_csv(f, sep="\t") for f in args.stats]

    doc = Document()
    doc.add_heading("Variant Classification Summary", level=1)

    # TODO: replace with your reordered ENG/ACVRL1/combined/v4/PALB2
    # table structure and sensitivity/specificity/PPV/NPV calculations
    for i, df in enumerate(classified_dfs):
        doc.add_heading(f"Table {i + 1}", level=2)
        table = doc.add_table(rows=1, cols=len(df.columns))
        for j, col in enumerate(df.columns):
            table.rows[0].cells[j].text = str(col)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)

    doc.save(args.output)


if __name__ == "__main__":
    main()
